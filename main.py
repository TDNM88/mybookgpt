import streamlit as st
import json
import os
from io import BytesIO
from groq import Groq
from fpdf import FPDF
from docx import Document

# Lấy API key từ môi trường hoặc Streamlit secrets
groq_api_key = st.secrets["groq"]["api_key"] if "groq" in st.secrets else os.getenv("GROQ_API_KEY")
client = Groq(api_key=groq_api_key)

# Khởi tạo session state nếu chưa có
if "book_outline" not in st.session_state:
    st.session_state["book_outline"] = None

if "book_content" not in st.session_state:
    st.session_state["book_content"] = ""
def read_uploaded_files(uploaded_files):
    combined_content = ""
    for file in uploaded_files:
        combined_content += file.read().decode("utf-8") + "\n"
    return combined_content

def generate_book_outline(topic_text, additional_instructions_prompt, model, groq_provider):
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": f"Tạo outline và tiêu đề cho cuốn sách về: {topic_text}\n{additional_instructions_prompt}"}],
            model=model,
            max_tokens=2000,
            temperature=0.7,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Lỗi khi tạo outline và tiêu đề sách: {str(e)}")
        return None

def generate_chapter_content(title, content, additional_instructions_prompt, model, groq_provider):
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": title + ": " + content + "\n" + additional_instructions_prompt}],
            model=model,
            max_tokens=2000,
            temperature=0.7,
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        st.error(f"Lỗi khi tạo nội dung chương: {str(e)}")
        return None
def save_to_docx(content, title="Cuốn Sách"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def save_to_pdf(content, title="Cuốn Sách"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, title)
    pdf.multi_cell(0, 10, content)
    file_stream = BytesIO()
    pdf.output(file_stream)
    file_stream.seek(0)
    return file_stream
# Streamlit UI
def main():
    st.set_page_config(page_title="AI Book Writer", layout="wide")

    st.title("AI Book Writer")
    st.write("Tạo sách tự động theo yêu cầu của bạn.")

    # Nhập thông tin từ người dùng
    book_topic = st.text_input("Chủ đề của cuốn sách:")
    writing_requirements = st.text_area("Yêu cầu về cách viết:")
    style_requirements = st.text_area("Yêu cầu về văn phong:")
    additional_instructions = st.text_area("Thông tin bổ sung:")

    uploaded_files = st.file_uploader("Tải lên các tệp văn bản", accept_multiple_files=True)
    reference_files = st.file_uploader("Tải lên các tệp tham khảo", accept_multiple_files=True)
    
    seed_content = read_uploaded_files(uploaded_files) if uploaded_files else ""
    reference_content = read_uploaded_files(reference_files) if reference_files else ""

    # Tạo outline và tiêu đề sách
    if st.button("Tạo Outline và Tiêu Đề Sách"):
        if len(book_topic) < 10:
            st.error("Chủ đề cuốn sách phải có ít nhất 10 ký tự.")
        else:
            additional_instructions_prompt = (
                additional_instructions
                + f"\nPhong cách viết: {style_requirements}\n"
                + f"Yêu cầu cách viết: {writing_requirements}\n"
            )
            if seed_content:
                additional_instructions_prompt += f"\nNội dung gợi ý: {seed_content}"
            if reference_content:
                additional_instructions_prompt += f"\nThông tin tham khảo: {reference_content}"

            # Tạo outline và tiêu đề sách
            st.session_state["book_outline"] = generate_book_outline(
                book_topic,
                additional_instructions_prompt,
                model="llama3-70b-8192",
                groq_provider=groq_api_key
            )
            if st.session_state["book_outline"]:
                st.write("### Outline và Tiêu Đề Sách Được Đề Xuất")
                st.text(st.session_state["book_outline"])

    # Tạo nội dung các chương
    if st.session_state["book_outline"] and st.button("Đồng ý với Outline và Tạo Nội Dung Các Chương"):
        st.write("### Đang tạo nội dung cho các chương...")
        book_outline_json = json.loads(st.session_state["book_outline"])
        for chapter in book_outline_json.keys():
            chapter_content = generate_chapter_content(
                title=chapter,
                content=book_outline_json[chapter],
                additional_instructions_prompt=additional_instructions_prompt,
                model="llama3-70b-8192",
                groq_provider=groq_api_key
            )
            if chapter_content:
                st.session_state["book_content"] += f"{chapter}\n\n{chapter_content}\n\n"
        
        if st.session_state["book_content"]:
            st.write("### Nội Dung Cuốn Sách:")
            st.text(st.session_state["book_content"])

            # Tải xuống nội dung cuốn sách
            if st.button("Tải Xuống Cuốn Sách"):
                txt_stream = StringIO(st.session_state["book_content"])
                docx_stream = save_to_docx(st.session_state["book_content"], title=book_topic)
                pdf_stream = save_to_pdf(st.session_state["book_content"], title=book_topic)

                st.download_button(
                    label="Tải xuống dưới dạng TXT",
                    data=txt_stream.getvalue(),
                    file_name="cuon_sach.txt",
                    mime="text/plain",
                )
                st.download_button(
                    label="Tải xuống dưới dạng DOCX",
                    data=docx_stream,
                    file_name="cuon_sach.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
                st.download_button(
                    label="Tải xuống dưới dạng PDF",
                    data=pdf_stream,
                    file_name="cuon_sach.pdf",
                    mime="application/pdf",
                )

if __name__ == "__main__":
    main()
